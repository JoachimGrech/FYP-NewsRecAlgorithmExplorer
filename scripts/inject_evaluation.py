"""
inject_evaluation.py  (v2)
Uses python-docx's proper paragraph insertion API via XML manipulation
with correct style IDs resolved from the document's style table.
"""

import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy

TEMPLATE_PATH = r"C:\Users\joach\Downloads\WordTemplateforFYPDissertations-01December2022(1).docx"
OUTPUT_PATH   = r"C:\Users\joach\Downloads\FYP_Dissertation_WithEvaluation.docx"

shutil.copy(TEMPLATE_PATH, OUTPUT_PATH)
doc = Document(OUTPUT_PATH)

# ── Resolve style IDs from the document ───────────────────────────────────────
def get_style_id(doc, style_name):
    for style in doc.styles:
        if style.name == style_name:
            return style.style_id
    raise ValueError(f"Style '{style_name}' not found in document.")

H2_ID         = get_style_id(doc, "Heading 2")
BODY_ID       = get_style_id(doc, "Dissertation Body")
BODY_FIRST_ID = get_style_id(doc, "Dissertation Body First Paragraph")

print(f"Style IDs: H2={H2_ID!r}, Body={BODY_ID!r}, BodyFirst={BODY_FIRST_ID!r}")

# ── Find Evaluation Heading 1 and next Heading 1 ─────────────────────────────
eval_h1_idx = None
next_h1_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 1" and "Evaluation" in p.text:
        eval_h1_idx = i
    elif eval_h1_idx is not None and p.style.name == "Heading 1":
        next_h1_idx = i
        break

print(f"Evaluation H1 at {eval_h1_idx}, next H1 at {next_h1_idx}")

# ── Remove placeholder paragraphs between the two H1s ────────────────────────
to_remove = [doc.paragraphs[i]._element for i in range(eval_h1_idx + 1, next_h1_idx)]
for elem in to_remove:
    elem.getparent().remove(elem)
print(f"Removed {len(to_remove)} placeholders.")

# ── Re-locate Evaluation H1 element ──────────────────────────────────────────
eval_anchor = None
for p in doc.paragraphs:
    if p.style.name == "Heading 1" and "Evaluation" in p.text:
        eval_anchor = p._element
        break

# ── Helper to build a paragraph XML element ───────────────────────────────────
def make_para(style_id, text):
    """Return a <w:p> lxml element with the given style and text."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)
    p.append(pPr)

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    return p

# ── Content ───────────────────────────────────────────────────────────────────
H2   = H2_ID
BODY = BODY_ID
BF   = BODY_FIRST_ID

sections = [
    (H2,   "Technical Evaluation Methodology"),
    (BF,   "The technical evaluation of the recommendation algorithms was conducted mechanically "
           "using custom Python evaluation scripts. The scripts simulated reading behaviour for "
           "diverse personas across the 5,000-article dataset. The objective was to quantitatively "
           "compare the baseline K-Means algorithm against the advanced Spherical Gaussian Mixture "
           "Model (GMM) under differing variables. Four primary metrics were evaluated: Average "
           "Relevance (cosine similarity), Topic Coverage (number of distinct clusters in the feed), "
           "Intra-List Similarity (ILS, measuring filter bubble strength), and Hit Rate."),

    (H2,   "Experiment 1: Scalability of Recommendations (Top-N Sweep)"),
    (BF,   "The first experiment evaluated algorithmic performance as the number of recommended "
           "articles (Top-N) increased from 5 to 25, with no artificial diversity injected."),
    (BODY, "Relevance: As the recommendation list size increased, both models exhibited a steady "
           "decline in average relevance as the system drew from progressively lower-ranked "
           "candidates. However, K-Means declined at approximately twice the rate of GMM. By the "
           "Top-25 mark, GMM maintained a relevance score of ~0.983, whereas K-Means fell to "
           "~0.915 \u2014 a gap of approximately 6.8 percentage points. This demonstrates that "
           "GMM\u2019s soft-clustering allows it to sustain high-quality recommendations for "
           "longer by drawing from overlapping topic boundaries."),
    (BODY, "[Figure 1: Comparative analysis of Average Relevance Score as the Top-N recommendation "
           "list size increases. The GMM algorithm maintains a higher, more stable relevance score "
           "(~0.98+) by utilizing soft-clustering, while the baseline K-Means algorithm exhibits "
           "a more pronounced steady decline.]"),
    (BODY, "Coverage: K-Means consistently restricted recommendations to approximately 1.25 topics "
           "regardless of list size \u2014 a near-flat line indicating a mathematically enforced "
           "single-topic filter bubble. The GMM, by contrast, naturally expanded its topic coverage "
           "from 2 topics at Top-5 to over 3.25 topics at Top-25, providing a genuinely "
           "multi-faceted feed without sacrificing relevance."),
    (BODY, "[Figure 2: Evaluation of Topic Coverage against Top-N list size. The GMM algorithm "
           "demonstrates superior topic diversity, expanding from 2 to 3.25 distinct topics as the "
           "list grows, whereas K-Means remains rigidly constrained to approximately 1.25 topics "
           "throughout, enforcing a filter bubble.]"),

    (H2,   "Experiment 2: Algorithmic Resilience to Diversity Injection"),
    (BF,   "The second experiment held the recommendation size at Top-10 and swept a stochastic "
           "diversity parameter from 0.0 to 0.4, injecting up to 40% noise to actively break "
           "filter bubbles."),
    (BODY, "Relevance vs Diversity: When forced to include diverse articles, K-Means relevance "
           "showed a steady decline, falling to ~0.89 at 40% noise. Remarkably, the GMM at its "
           "most diverse state (0.4) still maintained a higher relevance score (~0.96) than K-Means "
           "in its most focused state (0.0). This indicates that the GMM\u2019s understanding of "
           "overlapping semantic spaces allows it to \u201chide\u201d diversity within related "
           "topics, preserving the user experience while breaking filter bubbles."),
    (BODY, "[Figure 3: Algorithmic resilience to stochastic diversity injection. The GMM maintains "
           "a higher relevance floor even at high diversity (0.4) than K-Means does at zero "
           "diversity, proving its superior semantic robustness.]"),
    (BODY, "Filter Bubble Strength: Both models lowered Intra-List Similarity (ILS) as diversity "
           "was injected, indicating that the diversity mechanism is effective in both cases. "
           "Notably, GMM began from a slightly higher baseline ILS (~0.981 vs ~0.980 at 0.0 "
           "diversity) but exhibited a significantly steeper reduction, falling to ~0.931 at 40% "
           "noise compared to K-Means\u2019 ~0.952. This steeper descent indicates that GMM\u2019s "
           "probabilistic framework responds more dynamically to diversity injection \u2014 when "
           "pushed, it is architecturally more capable of genuinely diversifying its output."),
    (BODY, "[Figure 4: Intra-List Similarity (ILS) measured against increasing diversity injection. "
           "GMM exhibits a more significant reduction in filter bubble strength (lower ILS) compared "
           "to K-Means as diversity is introduced.]"),

    (H2,   "Experiment 3: Cluster Granularity Sensitivity (K-Sweep)"),
    (BF,   "A common technical critique of K-Means is that its rigidity can be mitigated by "
           "increasing the number of clusters (K). This experiment swept K across {5, 8, 15, 30, 50} "
           "to empirically test this hypothesis."),
    (BODY, "Results: Increasing K did not close the performance gap \u2014 it widened it. At K=5, "
           "both models performed similarly (~0.993 vs ~0.985, a 0.8% gap). However, as K increased, "
           "K-Means relevance degraded sharply, falling to ~0.830 at K=50, while GMM fell more "
           "gradually to ~0.879, producing a gap of nearly 5 percentage points. This reveals that "
           "K-Means becomes increasingly fragile at high granularity as articles are forced into "
           "hyper-specific clusters from which the recommender cannot escape."),
    (BODY, "An interesting trade-off was also observed in Hit Rate (categorical matching): at K=50, "
           "K-Means actually achieved a higher Hit Rate (~0.80 vs GMM\u2019s ~0.66) by locking onto "
           "the user\u2019s dominant topic label. However, this came at the direct cost of semantic "
           "depth \u2014 the recommendations matched the user\u2019s category but were less "
           "semantically coherent overall. GMM sacrifices rigid categorical matching in favour of "
           "genuine semantic proximity."),
    (BODY, "[Figure 5: Algorithmic sensitivity to cluster count (K). Both models exhibit declining "
           "relevance at higher K values, but K-Means degrades more rapidly, widening the performance "
           "gap from ~0.8% at K=5 to ~5% at K=50. GMM\u2019s probabilistic architecture proves more "
           "resilient to increased granularity.]"),

    (H2,   "Experiment 4: User Archetype Performance Analysis"),
    (BF,   "The final experiment evaluated performance across simulated user archetypes, ranging "
           "from \u201cHardcore Fans\u201d (single-topic focus) to \u201cEclectic Omnivores\u201d "
           "(multi-topic interests)."),
    (BODY, "Analysis: For single-topic users (e.g., \u201cHardcore Fan\u201d), K-Means performed "
           "competitively, achieving a relevance score virtually indistinguishable from GMM (~0.99 "
           "for both). However, for multi-topic archetypes \u2014 particularly \u201cEclectic\u201d "
           "(~0.93 vs ~0.98) and \u201cNews Junkie\u201d (~0.93 vs ~0.98) \u2014 GMM provided an "
           "absolute boost of approximately 5 percentage points. While modest in absolute terms, "
           "this represents a reduction in \u201crelevance error\u201d of nearly 70% (from 0.07 "
           "unexplained relevance gap to 0.02), confirming that the GMM architecture provides its "
           "strongest advantage for users whose interests span multiple intersecting topics."),
    (BODY, "[Figure 6: Average Relevance Score across distinct user behavioural archetypes. GMM "
           "shows significant performance gains over K-Means for multi-topic \u201cEclectic\u201d "
           "and \u201cNews Junkie\u201d readers.]"),

    (H2,   "Summary of Technical Findings"),
    (BF,   "The quantitative evaluation confirms the central hypothesis: a probabilistic Gaussian "
           "Mixture Model, when applied to SBERT embeddings, provides a recommendation engine that "
           "is more relevant, more diverse, and more resilient to filter bubbles than a traditional "
           "K-Means baseline. The soft-clustering approach effectively bridges the gap between "
           "mathematical similarity and human semantic nuance."),

    (H2,   "Proposed User Study Methodology"),
    (BF,   "Having mathematically validated the recommendation logic, the final phase of evaluation "
           "will involve a pedagogical user study to assess how effectively the visual radar charts "
           "communicate these algorithmic differences to non-expert users. Usability and educational "
           "effectiveness will be assessed through a small user study. Participants will be asked to "
           "complete predefined exploration tasks and provide feedback using questionnaires. Changes "
           "in user understanding before and after interaction with the system will be analysed."),
]

# Insert in reverse order so each paragraph ends up in correct sequence
current_anchor = eval_anchor
for style_id, text in sections:
    new_elem = make_para(style_id, text)
    current_anchor.addnext(new_elem)
    current_anchor = new_elem

doc.save(OUTPUT_PATH)
print(f"\nSaved: {OUTPUT_PATH}")

# ── Verify ────────────────────────────────────────────────────────────────────
doc2 = Document(OUTPUT_PATH)
printing = False
for p in doc2.paragraphs:
    if p.style.name == "Heading 1" and "Evaluation" in p.text:
        printing = True
    if p.style.name == "Heading 1" and "Conclusions" in p.text:
        break
    if printing:
        print(f"  [{p.style.name}] {p.text[:80]}")
