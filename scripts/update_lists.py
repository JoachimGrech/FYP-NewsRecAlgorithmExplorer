import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH   = r"C:\Users\joach\Downloads\FYP_Dissertation_Final_v2.docx"
NEW_OUTPUT_PATH = r"C:\Users\joach\Downloads\FYP_Dissertation_Final_v3.docx"

shutil.copy(OUTPUT_PATH, NEW_OUTPUT_PATH)
doc = Document(NEW_OUTPUT_PATH)

def make_para(style_name, text):
    style_id = None
    for s in doc.styles:
        if s.name == style_name:
            style_id = s.style_id
            break
            
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    if style_id:
        pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)
    p.append(pPr)

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    return p

figures = [
    "Figure 6.1 Comparative analysis of Average Relevance Score as the Top-N recommendation list size increases.",
    "Figure 6.2 Evaluation of Topic Coverage against Top-N list size.",
    "Figure 6.3 Algorithmic resilience to stochastic diversity injection.",
    "Figure 6.4 Intra-List Similarity (ILS) measured against increasing diversity injection.",
    "Figure 6.5 Algorithmic sensitivity to cluster count (K).",
    "Figure 6.6 Average Relevance Score across distinct user behavioural archetypes."
]

abbreviations = [
    "AI\tArtificial Intelligence",
    "BERT\tBidirectional Encoder Representations from Transformers",
    "CF\tCollaborative Filtering",
    "EU\tEuropean Union",
    "FYP\tFinal Year Project",
    "GMM\tGaussian Mixture Models",
    "ILS\tIntra-List Similarity",
    "LDA\tLatent Dirichlet Allocation",
    "SBERT\tSentence-BERT",
    "TF-IDF\tTerm Frequency-Inverse Document Frequency"
]

in_figures = False
in_tables = False
in_abbrevs = False

to_remove = []

for p in doc.paragraphs:
    if "List of Figures" in p.text and p.style.name.startswith("Headings"):
        in_figures = True
        in_tables = False
        in_abbrevs = False
        continue
    elif "List of Tables" in p.text and p.style.name.startswith("Headings"):
        in_figures = False
        in_tables = True
        in_abbrevs = False
        continue
    elif "List of Abbreviations" in p.text and p.style.name.startswith("Headings"):
        in_figures = False
        in_tables = False
        in_abbrevs = True
        continue
    elif p.style.name.startswith("Heading 1"):
        in_figures = False
        in_tables = False
        in_abbrevs = False
    
    if in_figures:
        to_remove.append(p)
    elif in_tables:
        to_remove.append(p)
    elif in_abbrevs:
        to_remove.append(p)

for p in to_remove:
    p._element.getparent().remove(p._element)

# Now inject the new lists
for p in doc.paragraphs:
    if "List of Figures" in p.text and p.style.name.startswith("Headings"):
        anchor = p._element
        for fig in figures:
            # Word uses tab stops for the page numbers, but without page numbers we just insert the text
            new_p = make_para("table of figures", fig)
            # add after
            anchor.addnext(new_p)
            anchor = new_p

for p in doc.paragraphs:
    if "List of Abbreviations" in p.text and p.style.name.startswith("Headings"):
        anchor = p._element
        # Remove any lingering "Note that..." paragraphs that might not have been caught if they had different styles, but to_remove caught them all
        for abbrev in abbreviations:
            new_p = make_para("Abbreviations", abbrev)
            anchor.addnext(new_p)
            anchor = new_p

doc.save(NEW_OUTPUT_PATH)
print("Lists updated.")
