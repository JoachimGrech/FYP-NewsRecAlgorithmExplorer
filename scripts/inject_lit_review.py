import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_PATH = r"C:\Users\joach\Downloads\FYP_Dissertation_WithEvaluation.docx"
OUTPUT_PATH   = r"C:\Users\joach\Downloads\FYP_Dissertation_Final.docx"

shutil.copy(TEMPLATE_PATH, OUTPUT_PATH)
doc = Document(OUTPUT_PATH)

def get_style_id(doc, style_name):
    for style in doc.styles:
        if style.name == style_name:
            return style.style_id
    raise ValueError(f"Style '{style_name}' not found.")

H1_ID = get_style_id(doc, "Heading 1")
H2_ID = get_style_id(doc, "Heading 2")
H3_ID = get_style_id(doc, "Heading 3")
BODY_ID = get_style_id(doc, "Dissertation Body")
BODY_FIRST_ID = get_style_id(doc, "Dissertation Body First Paragraph")

# Find Background H1 and Methodology H1
bg_idx = None
meth_idx = None

for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 1" and "Background" in p.text:
        bg_idx = i
    elif p.style.name == "Heading 1" and "Methodology" in p.text:
        meth_idx = i
        break

print(f"Background at {bg_idx}, Methodology at {meth_idx}")

# Delete everything between Background and Methodology
to_remove = [doc.paragraphs[i]._element for i in range(bg_idx, meth_idx)]
for elem in to_remove:
    elem.getparent().remove(elem)

# Locate the Methodology anchor
meth_anchor = None
for p in doc.paragraphs:
    if p.style.name == "Heading 1" and "Methodology" in p.text:
        meth_anchor = p._element
        break

def make_para(style_id, text):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)
    p.append(pPr)

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    return p

# Sections to insert:
sections = [
    (H1_ID, "Background"),
    (H2_ID, "The Role of Recommendation Algorithms in Modern News Consumption"),
    (BODY_FIRST_ID, "The rapid expansion of digital news has made algorithmic content filtering a central component of modern information consumption. Recommendation systems help users navigate vast volumes of online articles by predicting what content is most relevant to them. NVIDIA defines a recommendation system as a machine-learning model that uses user and item data “to help predict, narrow down, and find what people are looking for among an exponentially growing number of options” [4]."),
    (BODY_ID, "Platforms such as Google News, Facebook, and Apple News apply recommendation algorithms to personalise article feeds according to user preferences and behavioural patterns. This personalisation enhances convenience and engagement but simultaneously shapes how individuals encounter and interpret news [10]. As digital platforms increasingly mediate public discourse, recommendation systems now play a structural role in determining what users see, emphasising their relevance to both technological and societal contexts [5]."),
    
    (H2_ID, "Foundations of Recommendation Systems"),
    (BODY_FIRST_ID, "Recommendation systems traditionally employ a combination of content-based and collaborative filtering approaches, with many modern systems adopting hybrid methods to overcome limitations of individual techniques [9]. These methods provide the foundation upon which more advanced news recommendation systems are built."),
    
    (H3_ID, "Content-based Filtering"),
    (BODY_FIRST_ID, "Content-based filtering recommends items by matching the attributes of items a user has previously engaged with. Pazzani and Billsus characterise these systems as those that “analyze item descriptions to identify items that are of particular interest to the user” [8]."),
    (BODY_ID, "In the context of news, content-based systems rely heavily on text analysis techniques, including keyword extraction, Term Frequency, Inverse Document Frequency vectors, topic modelling, and contextual embeddings. Topic modelling methods such as Latent Dirichlet Allocation (LDA) [12] identify latent thematic structures within text, while models like BERT [13] provide deep semantic representations that capture contextual meaning."),
    (BODY_ID, "Content-based systems avoid cold-start problems on the user side and provide transparent item-feature relationships; however, they risk overspecialisation by recommending only narrowly similar articles."),
    
    (H3_ID, "Collaborative Filtering"),
    (BODY_FIRST_ID, "Collaborative filtering (CF) leverages patterns in user–item interactions to generate predictions. User-user CF assumes that users with similar historical preferences will favour similar items, while item-item CF compares items based on shared audience engagement [6]."),
    (BODY_ID, "Matrix factorisation techniques, presented by Koren et al. [7], significantly improved CF by representing users and items in a shared latent space, enabling scalable and accurate predictions. CF remains a core technique in commercial recommendation platforms but suffers from cold-start limitations for new users and new items."),
    (BODY_ID, "In news contexts, CF is challenged further by rapidly changing content and short article lifecycles, making hybridisation with content-based methods especially valuable."),
    
    (H3_ID, "Advanced Content Representation and Soft Clustering"),
    (BODY_FIRST_ID, "Traditional content-based recommendation systems often rely on \"hard clustering\" techniques, such as K-Means, which strictly assign an article to a single, mutually exclusive topic. While computationally efficient, this approach struggles to accurately represent news content, which is inherently multi-dimensional. For example, an article discussing AI regulations intersects both \"Technology\" and \"Politics\". Forcing such articles into a single category results in a loss of semantic nuance and limits the accuracy of content-based recommendations."),
    
    (H3_ID, "Probabilistic Topic Modelling via Gaussian Mixture Models (GMM)"),
    (BODY_FIRST_ID, "To address the limitations of hard boundaries, advanced systems increasingly employ \"soft clustering\" methods like Gaussian Mixture Models (GMM). A GMM is a probabilistic model that assumes data points are generated from a mixture of a finite number of Gaussian distributions [27]. Instead of a binary topic assignment, GMM calculates posterior probabilities—a probability distribution indicating an article's likelihood of belonging to each cluster. This allows the system to understand that an article is, for example, 60% Topic A and 40% Topic B, enabling far more accurate content matching and representation [27]."),
    
    (H3_ID, "Synergising Deep Semantic Embeddings with Probabilistic Clustering"),
    (BODY_FIRST_ID, "Modern recommendation architectures frequently combine probabilistic clustering with deep semantic embeddings to create robust representations of textual content. Sentence-BERT (SBERT), a modification of the pre-trained BERT network, generates semantically meaningful, dense vector representations of sentences or documents [25]. When SBERT embeddings are processed using clustering algorithms, the system can capture deep contextual meanings that traditional frequency-based models (like TF-IDF or LDA) miss [26]. By applying GMM over these dense semantic vectors, recommender systems can map complex text into overlapping topic distributions, significantly improving both the accuracy and explainability of the resulting recommendations."),
    
    (H1_ID, "Literature Review"),
    (H2_ID, "Algorithmic Influence on News Exposure and Information Access"),
    (BODY_FIRST_ID, "Recommendation algorithms strongly influence how individuals encounter information. By filtering news based on predicted relevance, algorithms can increase personalisation but also limit exposure to diverse viewpoints."),
    (BODY_ID, "Das et al. demonstrated how personalised large-scale systems such as Google News adapt article selection based on user behaviour [10]. Liu et al. further showed that real-time click behaviour can meaningfully alter users’ news feeds [11]."),
    (BODY_ID, "While these systems enhance usability, they have a recognised impact on information access patterns, often promoting content aligned with prior behaviour. Studies indicate that algorithmic filtering can contribute to ideological segregation when users are repeatedly exposed to content confirming pre-existing preferences [15]."),
    (BODY_ID, "Thus, recommendation algorithms not only distribute information but also actively shape news consumption patterns and public understanding."),
    
    (H2_ID, "Public Understanding and Awareness of Algorithmic Personalisation"),
    (BODY_FIRST_ID, "Despite their ubiquity, news recommendation algorithms often remain poorly understood by the general public. Gillespie argues that algorithms have become powerful “gatekeepers” of information, yet most users remain unaware of how or why certain content appears in their feeds [21]."),
    (BODY_ID, "Research shows that people tend to underestimate the degree of personalisation occurring on digital platforms and often do not recognise that the news they receive is algorithmically curated [19]."),
    (BODY_ID, "A lack of awareness can reduce users’ ability to critically evaluate the content presented to them and undermines their capacity to identify bias, manipulation, or skewed information exposure. As a result, improving public understanding of algorithmic personalisation has become a priority within digital literacy and media education initiatives."),
    
    (H2_ID, "The Black-box Nature of News Recommendation Algorithms"),
    (BODY_FIRST_ID, "One of the primary challenges in recommendation systems is their opacity. Many commercial recommenders operate as black boxes, revealing little about the parameters, data, or logic used to produce recommendations."),
    (BODY_ID, "Zhang and Chen argue that opaque systems hinder user trust and accountability, making it difficult for individuals to assess why particular items are shown [18]. Similarly, Burke and later scholars highlight how hybridised and deep-learning-based systems, now common in news platforms, further obscure explainability due to their complexity [9]."),
    (BODY_ID, "This opacity contributes to issues such as algorithmic bias, filter bubbles, and unequal representation. Pariser popularised the notion of the filter bubble, describing how opaque personalisation processes isolate users within narrow informational domains [14]."),
    (BODY_ID, "The black-box nature of news recommenders thus raises significant concerns related to transparency, fairness, and democratic accountability."),
    
    (H2_ID, "Explainability in Recommendation Systems"),
    (BODY_FIRST_ID, "Explainability has become an increasingly important area of research within recommender systems, particularly as algorithmic systems play a greater role in shaping how users access digital information. Traditional recommendation algorithms, especially those based on complex machine learning or deep learning models, often operate as opaque systems whose internal decision processes are difficult for users to interpret. This lack of transparency can make it challenging for users to understand why particular items are recommended, reinforcing the perception that recommendation systems function as “black boxes” [18]."),
    (BODY_ID, "Explainable recommender systems aim to address this limitation by providing understandable reasons for algorithmic outputs. Rather than presenting recommendations without context, these systems attempt to communicate the factors that influenced the recommendation process. Zhang and Chen describe explainable recommendation methods as approaches that provide explanations to make recommendation results more transparent and interpretable to users [18]. By clarifying how recommendations are generated, explainability can improve user trust, increase perceived reliability, and encourage more informed interaction with algorithmic systems [28]."),
    (BODY_ID, "Explainability also provides important benefits from a system development perspective. Transparent recommendation mechanisms allow researchers and developers to analyse how algorithms behave, detect biases in the data, and better understand how different inputs influence the final recommendations. Tintarev and Masthoff emphasise that explanations can support both system evaluation and user satisfaction by improving transparency and helping users understand the reasoning behind algorithmic decisions [22]."),
    (BODY_ID, "Several approaches have been developed to generate explanations within recommender systems. One common technique involves feature-based explanations, which highlight the attributes of items that contributed to the recommendation. In content-based recommender systems, recommendations can be explained by identifying similarities between the recommended item and the user’s previous interactions. For example, if a user frequently reads articles related to a particular topic, the system may explain that a recommendation was generated because the article contains similar keywords or themes."),
    (BODY_ID, "Another approach involves collaborative explanations that rely on behavioural patterns observed among similar users. In collaborative filtering systems, recommendations can be justified by referencing the behaviour of users with comparable preferences. For instance, a system may indicate that an article is recommended because it was frequently read or positively rated by users with similar reading histories. Early work by Herlocker, Konstan, and Riedl demonstrated how collaborative filtering systems can generate explanations by highlighting relationships between user preferences and recommendation outcomes [23]. Furthermore, as modern platforms increasingly adopt hybrid approaches, recent studies have demonstrated that providing personalised explanations, tailoring the explanation style to the individual user, can significantly enhance the effectiveness and transparency of hybrid recommender systems [29]."),
    (BODY_ID, "More advanced explainability techniques have also been proposed to interpret complex machine-learning models used in modern recommender systems. Methods such as influence functions and feature importance analysis aim to identify which elements of the input data most strongly affect the prediction produced by the model. These techniques provide insight into the behaviour of otherwise opaque algorithms and help researchers understand how specific training data points influence recommendation outputs [24]."),
    (BODY_ID, "Explainability is particularly important in the context of news recommendation systems. Because algorithmic recommenders influence which news articles users encounter, the transparency of these systems has implications for information exposure and public discourse. When users are unaware of how recommendations are generated, they may underestimate the degree to which algorithmic personalisation shapes their information environment. Providing explanations for recommended content can therefore help users recognise how their previous behaviour, preferences, or interactions influence the news they receive."),
    (BODY_ID, "For these reasons, explainability has become an important design consideration for modern recommender systems. By making algorithmic decisions more transparent and interpretable, explainable recommendation systems can improve user trust, support critical engagement with personalised content, and contribute to greater accountability in algorithmic information filtering. These principles are particularly relevant for educational systems that aim to demonstrate how recommendation algorithms operate, such as the interactive news recommendation algorithm explorer proposed in this project."),
    
    (H2_ID, "Educational and Interactive Approaches to Algorithm Transparency"),
    (BODY_FIRST_ID, "Interactive tools and visualisations have emerged as effective means of demystifying algorithmic processes. Systems such as TensorFlow Playground show how real-time model adjustments can help users understand model behaviour [16]. Influence-function-based interactive tools demonstrate how specific data points affect predictions, offering insight into algorithmic decision-making [24]."),
    (BODY_ID, "In the domain of recommenders, Ekstrand et al.’s LensKit framework supports transparent experimentation with recommendation algorithms, enabling users to explore algorithm behaviour and parameter effects [17]."),
    (BODY_ID, "Interactive explainability approaches, highlighted by Zhang and Chen [18], can help users build accurate mental models of system behaviour."),
    (BODY_ID, "These precedents demonstrate the value of educational interfaces, such as the proposed Interactive News Recommendation Algorithm Explorer, in improving public comprehension of algorithmic personalisation."),
    
    (H2_ID, "Ethical Considerations in Simulated News Recommendation Systems"),
    (BODY_FIRST_ID, "Simulating news recommendation systems for educational purposes requires careful ethical consideration. O’Neil warns that algorithmic systems can reinforce inequality and amplify harmful outcomes when deployed without safeguards [19]. In the context of news, these harms may manifest as sensationalism, emotional manipulation, or the reinforcement of extreme viewpoints."),
    (BODY_ID, "Tufekci argues that automated curation systems have the power to influence public perception and behaviour, often in ways users do not recognise [20]."),
    (BODY_ID, "Simulated systems must therefore:"),
    (BODY_ID, "•	Avoid replicating harmful biases."),
    (BODY_ID, "•	Clearly communicate that they are educational models, not real-world predictors."),
    (BODY_ID, "•	Present diverse and balanced news sources, and avoid oversimplifying complex algorithmic behaviour."),
    (BODY_ID, "Ensuring transparency, fairness, and ethical integrity in such systems is crucial to promoting meaningful public understanding rather than contributing to misinformation or misunderstandings about real-world algorithms.")
]

# We must insert BEFORE meth_anchor. We can do this by inserting sequentially before meth_anchor.
# Wait, meth_anchor.addprevious(elem) is available in lxml.
for style_id, text in sections:
    new_elem = make_para(style_id, text)
    meth_anchor.addprevious(new_elem)

doc.save(OUTPUT_PATH)
print(f"Saved successfully to {OUTPUT_PATH}")
